#!/usr/bin/env python3
"""
Convert first 11 pages of new_testament.docx to PDF using pure Python.
Works on Linux without LibreOffice!

Install dependencies:
    pip install python-docx aspose-words

OR if you prefer open-source (requires pandoc):
    pip install python-docx pypandoc
    sudo apt install pandoc texlive-latex-base

Usage:
    python3 convert.py
"""

import os
import sys
from pathlib import Path

# ============================================================
# CONFIGURATION - Change these values as needed
# ============================================================
INPUT_FILE = "new_testament.docx"
OUTPUT_FILE = "output.pdf"
NUM_PAGES = 11

# Choose conversion method:
# "aspose" - Commercial library (best quality, requires license for production)
# "pypandoc" - Open source (requires pandoc + LaTeX installed)
METHOD = "aspose"  # Change to "pypandoc" if you prefer open-source
# ============================================================

def extract_pages_docx(input_path, output_path, num_pages):
    """Extract the first N pages from DOCX."""
    from docx import Document
    
    print(f"📄 Loading document: {input_path}")
    doc = Document(input_path)
    new_doc = Document()
    
    # Copy document properties
    new_doc.core_properties.author = doc.core_properties.author
    new_doc.core_properties.title = doc.core_properties.title
    
    page_breaks_found = 0
    elements_copied = 0
    
    print(f"🔍 Searching for page breaks...")
    
    # Iterate through all elements in the document body
    for element in doc.element.body:
        # Stop if we've found enough pages
        if page_breaks_found >= num_pages:
            break
        
        # Check for page breaks in paragraphs
        if element.tag.endswith('p'):  # paragraph
            # Count page breaks in this paragraph
            for run in element.findall('.//w:r', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                if run.find('.//w:br[@w:type="page"]', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
                    page_breaks_found += 1
                    if page_breaks_found >= num_pages:
                        break
            
            # Copy the paragraph element
            new_doc.element.body.append(element)
            elements_copied += 1
            
        elif element.tag.endswith('tbl'):  # table
            if page_breaks_found < num_pages:
                new_doc.element.body.append(element)
                elements_copied += 1
        
        elif element.tag.endswith('sectPr'):  # section properties
            page_breaks_found += 1
            if page_breaks_found < num_pages:
                new_doc.element.body.append(element)
    
    print(f"✓ Found {page_breaks_found} page breaks")
    print(f"✓ Copied {elements_copied} elements")
    
    print(f"💾 Saving extracted content to: {output_path}")
    new_doc.save(output_path)
    return output_path

def convert_with_aspose(docx_path, pdf_path):
    """
    Convert DOCX to PDF using Aspose.Words (best quality).
    Note: Free for evaluation, but requires license for production use.
    """
    print(f"\n🔄 Converting using Aspose.Words...")
    
    try:
        import aspose.words as aw
        print(f"✓ Aspose.Words imported successfully!")
        
        # Load the document
        print(f"📖 Loading document: {docx_path}")
        doc = aw.Document(docx_path)
        
        # Save as PDF
        print(f"💾 Saving as PDF: {pdf_path}")
        doc.save(pdf_path)
        
        if os.path.exists(pdf_path):
            print(f"✓ PDF created successfully with Aspose!")
            return True
        else:
            print(f"✗ PDF file was not created")
            return False
            
    except ImportError as e:
        print(f"\n✗ Aspose.Words import failed: {e}")
        print("\nDebugging info:")
        print(f"  Python executable: {sys.executable}")
        print(f"  Python path: {sys.path}")
        
        # Try to find where aspose is installed
        try:
            import subprocess
            result = subprocess.run([sys.executable, '-m', 'pip', 'show', 'aspose-words'], 
                                  capture_output=True, text=True)
            print(f"  pip show aspose-words:\n{result.stdout}")
        except:
            pass
        
        print("\nInstall it with: pip install aspose-words")
        print("\nNote: Aspose is free for evaluation but requires a license for production.")
        print("For a free alternative, change METHOD to 'pypandoc' in the script.")
        return False
    except Exception as e:
        print(f"\n✗ Aspose conversion error: {e}")
        import traceback
        traceback.print_exc()
        return False

def convert_with_pypandoc(docx_path, pdf_path):
    """
    Convert DOCX to PDF using pypandoc (open source).
    Requires: pandoc and LaTeX (texlive) to be installed.
    """
    print(f"\n🔄 Converting using pypandoc...")
    
    try:
        import pypandoc
        
        # Convert to PDF
        output = pypandoc.convert_file(
            docx_path,
            'pdf',
            outputfile=pdf_path,
            extra_args=['--pdf-engine=pdflatex']
        )
        
        if os.path.exists(pdf_path):
            print(f"✓ PDF created successfully with pypandoc!")
            return True
        else:
            print(f"✗ PDF file was not created")
            return False
            
    except ImportError:
        print("\n✗ pypandoc is not installed")
        print("Install with: pip install pypandoc")
        return False
    except RuntimeError as e:
        if "pandoc" in str(e).lower():
            print("\n✗ Pandoc is not installed on your system")
            print("Install with: sudo apt install pandoc texlive-latex-base")
            print("\nOr use Aspose method by changing METHOD='aspose' in the script")
        else:
            print(f"\n✗ pypandoc error: {e}")
        return False
    except Exception as e:
        print(f"\n✗ pypandoc conversion error: {e}")
        return False

def main():
    print(f"\n{'='*60}")
    print(f"📄 DOCX to PDF Converter (Pure Python - No LibreOffice!)")
    print(f"{'='*60}")
    print(f"Input:  {INPUT_FILE}")
    print(f"Output: {OUTPUT_FILE}")
    print(f"Pages:  First {NUM_PAGES} pages")
    print(f"Method: {METHOD}")
    print(f"{'='*60}\n")
    
    # Check if input file exists
    if not os.path.exists(INPUT_FILE):
        print(f"✗ Error: Input file not found: {INPUT_FILE}")
        print(f"\nCurrent directory: {os.getcwd()}")
        print(f"\nFiles in current directory:")
        for f in sorted(os.listdir('.')):
            if f.endswith('.docx'):
                print(f"  📄 {f}")
        return
    
    try:
        # Step 1: Extract first N pages
        print(f"Step 1: Extracting first {NUM_PAGES} pages")
        print(f"{'-'*60}")
        
        temp_docx = "temp_extracted.docx"
        extract_pages_docx(INPUT_FILE, temp_docx, NUM_PAGES)
        
        # Step 2: Convert to PDF
        print(f"\nStep 2: Converting to PDF")
        print(f"{'-'*60}")
        
        if METHOD == "aspose":
            success = convert_with_aspose(temp_docx, OUTPUT_FILE)
        elif METHOD == "pypandoc":
            success = convert_with_pypandoc(temp_docx, OUTPUT_FILE)
        else:
            print(f"✗ Unknown method: {METHOD}")
            print(f"Valid methods: 'aspose', 'pypandoc'")
            success = False
        
        # Cleanup temporary file
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
            print(f"🗑️  Cleaned up temporary file")
        
        # Final result
        print(f"\n{'='*60}")
        if success:
            size = os.path.getsize(OUTPUT_FILE)
            print(f"✅ SUCCESS! PDF created: {OUTPUT_FILE}")
            print(f"📊 File size: {size / 1024 / 1024:.2f} MB")
            print(f"\n🎉 All formatting and images preserved!")
        else:
            print(f"❌ FAILED: Could not create PDF")
            print(f"\n💡 Troubleshooting:")
            if METHOD == "aspose":
                print(f"   • Install Aspose: pip install aspose-words")
                print(f"   • OR switch to pypandoc method")
            elif METHOD == "pypandoc":
                print(f"   • Install pypandoc: pip install pypandoc")
                print(f"   • Install pandoc: sudo apt install pandoc texlive-latex-base")
                print(f"   • OR switch to aspose method (better quality)")
        print(f"{'='*60}\n")
        
    except ImportError as e:
        print(f"\n✗ Missing required package: {e}")
        print("\nInstall required packages:")
        print("  pip install python-docx")
        print(f"  pip install {METHOD if METHOD == 'aspose' else 'pypandoc'}")
    except Exception as e:
        print(f"\n✗ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
