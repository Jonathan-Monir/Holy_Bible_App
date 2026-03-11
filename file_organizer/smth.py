# Requires: pip install python-docx
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

def read_first_n_paragraphs(path, n=4):
    doc = Document(path)
    return [p.text for p in doc.paragraphs[:n]]

def read_section_footers(path):
    doc = Document(path)
    footers = []
    for i, sec in enumerate(doc.sections):
        text = '\n'.join(p.text for p in sec.footer.paragraphs if p.text)
        if text:
            footers.append((i, text))
    return footers

def read_footnotes(path):
    footnotes = []
    try:
        with zipfile.ZipFile(path) as z:
            data = z.read('word/footnotes.xml')
    except KeyError:
        return footnotes  # no footnotes part
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    root = ET.fromstring(data)
    for fn in root.findall('w:footnote', ns):
        # skip separator/continuationSeparator entries
        typ = fn.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
        if typ in ('separator', 'continuationSeparator'):
            continue
        parts = [t.text or '' for t in fn.findall('.//w:t', ns)]
        footnotes.append(''.join(parts))
    return footnotes

if __name__ == '__main__':
    path = 'old_testament.docx'
    print("First 4 paragraphs:")
    for i, p in enumerate(read_first_n_paragraphs(path, 4), 1):
        print(f"{i}: {p}")

    print("\nSection footers (if any):")
    for sec_idx, text in read_section_footers(path):
        print(f"section {sec_idx}: {text}")

    print("\nFootnotes (if any):")
    for i, fn in enumerate(read_footnotes(path), 1):
        print(f"{i}: {fn}")
